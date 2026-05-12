"""File Organizer — sort a folder's files into category subfolders.

Single-file Tk application. Categorization combines filename patterns with
optional PDF/DOCX text inspection for CV detection. Built to be packaged as a
standalone Windows executable via PyInstaller --onefile.
"""
from __future__ import annotations

import functools
import io
import json
import logging
import os
import queue
import re
import shutil
import ssl
import subprocess
import sys
import threading
import time
import tkinter as tk
import unicodedata
import urllib.request
import webbrowser
from datetime import datetime
from pathlib import Path
from tkinter import END, filedialog, messagebox, ttk

APP_NAME = "File Organizer"
APP_VERSION = "1.4.0"
UNDO_LOG_NAME = ".file-organizer-undo.json"
BMC_URL = "https://buymeacoffee.com/bturksoy"
UPDATE_API_URL = (
    "https://api.github.com/repos/bturksoy/auto-file-organizer/releases/latest"
)
UPDATE_HTTP_TIMEOUT = 8
DEFAULT_LANG = "en"
RECENT_FOLDERS_LIMIT = 8
MIN_AUTO_INTERVAL_MIN = 1


# ---------------------------------------------------------------------------
# Resource loading
#
# Translations and rule data live in JSON files under resources/ so they can
# be edited (and contributed) without touching Python. At runtime we load
# them from either the dev tree or the PyInstaller _MEIPASS extraction dir.
# ---------------------------------------------------------------------------

def _resources_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS) / "resources"
    return Path(__file__).resolve().parent / "resources"


def _read_json(*parts) -> object:
    path = _resources_dir().joinpath(*parts)
    return json.loads(path.read_text(encoding="utf-8"))


def _load_i18n_bundles() -> tuple[dict, dict, dict]:
    """Walk resources/i18n/*.json and assemble the three runtime tables.

    Returns (languages, ui_strings, category_names) keyed by language code.
    Missing or malformed files are skipped — a broken translation should not
    take the whole app down.
    """
    languages: dict[str, str] = {}
    ui: dict[str, dict] = {}
    cats: dict[str, dict] = {}
    i18n_dir = _resources_dir() / "i18n"
    if not i18n_dir.is_dir():
        return languages, ui, cats
    for path in sorted(i18n_dir.glob("*.json")):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        meta = data.get("_meta") or {}
        code = meta.get("code") or path.stem
        languages[code] = meta.get("name") or code.upper()
        ui[code] = data.get("strings", {})
        cats[code] = data.get("categories", {})
    return languages, ui, cats


def _load_extensions() -> dict[str, set[str]]:
    raw = _read_json("data", "extensions.json")
    return {key: set(exts) for key, exts in raw.items() if not key.startswith("_")}


def _load_cv_keywords() -> tuple[tuple[str, ...], tuple[str, ...]]:
    raw = _read_json("data", "cv_keywords.json")
    strong = tuple(raw.get("strong", []))
    weak = tuple(raw.get("weak", []))
    return strong, weak


def _load_skip_names() -> set[str]:
    raw = _read_json("data", "skip_names.json")
    if isinstance(raw, list):
        return set(raw)
    return set(raw.get("names", []))

# pypdf is noisy about malformed streams. We don't want to crash on it.
logging.getLogger("pypdf").setLevel(logging.ERROR)


# ---------------------------------------------------------------------------
# Settings persistence
# ---------------------------------------------------------------------------

def _config_path() -> Path:
    base = os.environ.get("APPDATA") or str(Path.home() / ".config")
    return Path(base) / "FileOrganizer" / "settings.json"


def load_settings() -> dict:
    path = _config_path()
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


def save_settings(data: dict) -> None:
    path = _config_path()
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8"
        )
    except OSError:
        # If we can't write settings, the app should still run.
        pass


# ---------------------------------------------------------------------------
# Internationalization
# ---------------------------------------------------------------------------

LANGUAGES, UI_STRINGS, CATEGORY_NAMES = _load_i18n_bundles()
if DEFAULT_LANG not in LANGUAGES:
    # Defensive: ensure we always have an English fallback even if the
    # bundled JSON is missing or unreadable. Tk messageboxes don't really
    # work pre-mainloop, so we just print a hint and continue with an
    # empty dict — the I18N class will then echo keys back as text.
    print("WARNING: no English translation bundle found at resources/i18n/",
          file=sys.stderr)
    LANGUAGES.setdefault(DEFAULT_LANG, "English")
    UI_STRINGS.setdefault(DEFAULT_LANG, {})
    CATEGORY_NAMES.setdefault(DEFAULT_LANG, {})


class I18N:
    """Tiny translation lookup with a default-language fallback."""

    def __init__(self, lang: str = DEFAULT_LANG) -> None:
        self._lang = lang if lang in UI_STRINGS else DEFAULT_LANG

    @property
    def lang(self) -> str:
        return self._lang

    def set_language(self, lang: str) -> None:
        if lang in UI_STRINGS:
            self._lang = lang

    def t(self, key: str, **kwargs) -> str:
        table = UI_STRINGS.get(self._lang, UI_STRINGS[DEFAULT_LANG])
        text = table.get(key) or UI_STRINGS[DEFAULT_LANG].get(key, key)
        return text.format(**kwargs) if kwargs else text


i18n = I18N()


# ---------------------------------------------------------------------------
# Update checking
# ---------------------------------------------------------------------------

def _version_tuple(v: str) -> tuple:
    """Convert a version string like 'v1.10.2-beta' into (1, 10, 2)."""
    parts = []
    for chunk in v.lstrip("v").split("."):
        digits = re.match(r"\d+", chunk)
        parts.append(int(digits.group()) if digits else 0)
    return tuple(parts) or (0,)


def is_newer_version(remote: str, local: str) -> bool:
    return _version_tuple(remote) > _version_tuple(local)


def _human_size(n: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.1f} {unit}" if unit != "B" else f"{n} B"
        n /= 1024
    return f"{n:.1f} TB"


def _ssl_context() -> ssl.SSLContext:
    """Build an SSL context that works across diverse Windows setups.

    Order of preference:
      1. truststore — uses the Windows native cert store, which is the
         only thing that knows about corporate root CAs installed by IT
         policy (common cause of "unable to get local issuer certificate").
      2. certifi — bundled CA list, works on most consumer machines.
      3. Python default — last resort.
    """
    try:
        import truststore
        return truststore.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
    except Exception:
        pass
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except Exception:
        pass
    return ssl.create_default_context()


def fetch_latest_release_info() -> dict | None:
    """Query GitHub for the latest release. None on any failure.

    The returned dict contains: version, url (exe asset), size, page (html_url).
    """
    try:
        req = urllib.request.Request(
            UPDATE_API_URL,
            headers={
                "User-Agent": f"FileOrganizer/{APP_VERSION}",
                "Accept": "application/vnd.github+json",
            },
        )
        with urllib.request.urlopen(req, timeout=UPDATE_HTTP_TIMEOUT,
                                    context=_ssl_context()) as r:
            data = json.load(r)
    except Exception:
        return None

    tag = (data.get("tag_name") or "").strip()
    if not tag:
        return None

    exe_asset = next(
        (a for a in data.get("assets", [])
         if a.get("name", "").lower().endswith(".exe")),
        None,
    )
    if not exe_asset:
        return None

    return {
        "version": tag.lstrip("v"),
        "url": exe_asset["browser_download_url"],
        "size": exe_asset.get("size", 0),
        "page": data.get("html_url"),
    }


def is_running_frozen() -> bool:
    """True when launched as a PyInstaller-bundled exe."""
    return getattr(sys, "frozen", False)


def apply_update(asset_url: str, on_progress=None) -> None:
    """Download the new exe and spawn a swap-and-restart helper.

    Does NOT terminate the current process. The caller is responsible for
    shutting down the app cleanly so the running exe releases its file lock;
    only then can the helper move the new binary into place.
    """
    if not is_running_frozen():
        raise RuntimeError("Auto-update only works from the packaged exe.")

    current = Path(sys.executable).resolve()
    new_path = current.with_name(current.stem + ".new.exe")

    # Stream the download so we can report progress.
    req = urllib.request.Request(
        asset_url,
        headers={"User-Agent": f"FileOrganizer/{APP_VERSION}"},
    )
    with urllib.request.urlopen(req, timeout=UPDATE_HTTP_TIMEOUT * 4,
                                context=_ssl_context()) as resp:
        total = int(resp.headers.get("Content-Length") or 0)
        done = 0
        with open(new_path, "wb") as out:
            while True:
                chunk = resp.read(65536)
                if not chunk:
                    break
                out.write(chunk)
                done += len(chunk)
                if on_progress and total:
                    on_progress(done, total)

    # Swap-and-restart batch.
    # - Initial 3s grace so the parent has time to fully exit
    # - Bounded retry (~40s total) so we never loop forever if the file
    #   somehow stays locked or the new binary disappears
    # - All output suppressed; only one cmd is ever spawned (detached)
    bat = current.with_name("_fo_update.bat")
    script = (
        '@echo off\r\n'
        'setlocal\r\n'
        'set /a RETRIES=40\r\n'
        'ping 127.0.0.1 -n 4 >nul 2>&1\r\n'
        ':loop\r\n'
        f'move /y "{new_path}" "{current}" >nul 2>&1\r\n'
        f'if not exist "{new_path}" goto done\r\n'
        'set /a RETRIES=RETRIES-1\r\n'
        'if %RETRIES% LEQ 0 goto done\r\n'
        'ping 127.0.0.1 -n 2 >nul 2>&1\r\n'
        'goto loop\r\n'
        ':done\r\n'
        f'if exist "{current}" start "" "{current}"\r\n'
        '(goto) 2>nul & del "%~f0"\r\n'
    )
    bat.write_text(script, encoding="ascii")

    # CREATE_NO_WINDOW alone gives us a hidden, detached child cmd.
    # (DETACHED_PROCESS + CREATE_NO_WINDOW together can fail on some Windows
    # builds; sticking to the simpler flag is more reliable.)
    subprocess.Popen(
        ["cmd.exe", "/c", str(bat)],
        creationflags=0x08000000,
        close_fds=True,
        stdin=subprocess.DEVNULL,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )



# Category display names come from resources/i18n/<lang>.json.
def category_display(key: str, lang: str | None = None) -> str:
    """Localized folder/display name for an internal category key."""
    lang = lang or i18n.lang
    return CATEGORY_NAMES.get(lang, CATEGORY_NAMES["en"]).get(key, key)


def all_category_names_for_lang(lang: str) -> set[str]:
    return set(CATEGORY_NAMES.get(lang, CATEGORY_NAMES["en"]).values())


def all_known_category_names() -> set[str]:
    """Every localized name across every supported language.

    Used when walking a folder so we never recurse into a previously created
    category subfolder, regardless of the language it was created in.
    """
    names: set[str] = set()
    for lang in CATEGORY_NAMES:
        names |= all_category_names_for_lang(lang)
    return names


# ---------------------------------------------------------------------------
# Text normalization and CV-content detection
# ---------------------------------------------------------------------------

def _normalize(s: str) -> str:
    """Lowercase, strip diacritics, fold dotless ı to i.

    The dotless ı (U+0131) is a standalone letter rather than a combining
    sequence, so NFD does not decompose it. We fold it explicitly so that
    Turkish filenames like 'Kılavuz' match the keyword 'kilavuz'.
    """
    s = unicodedata.normalize("NFD", s.casefold())
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return s.replace("ı", "i")


def _maybe_despace(text: str) -> str:
    """Repair PDF text that was rendered character-by-character.

    Some PDFs (commonly those exported from Canva/Figma) place each glyph as
    its own text object. Most extractors then emit one space between every
    character: "W o r k  E x p e r i e n c e". We detect that pattern via
    the ratio of single-character tokens and collapse the inter-character
    spaces while preserving word breaks (double spaces).
    """
    tokens = text.split()
    if len(tokens) < 10:
        return text
    single = sum(1 for tok in tokens if len(tok) == 1) / len(tokens)
    if single < 0.4:
        return text
    placeholder = "\x00"
    out = re.sub(r"  +", placeholder, text)
    out = out.replace(" ", "")
    return out.replace(placeholder, " ")


# CV detection keyword lists come from resources/data/cv_keywords.json.
_CV_STRONG_RAW, _CV_WEAK_RAW = _load_cv_keywords()
_CV_STRONG = tuple(_normalize(k) for k in _CV_STRONG_RAW)
_CV_WEAK = tuple(_normalize(k) for k in _CV_WEAK_RAW)


def _read_pdf_text(path: Path, max_pages: int = 4) -> str:
    """Extract text from the first few pages. Returns '' on any failure."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path), strict=False)
        if reader.is_encrypted:
            return ""
        parts = []
        for page in reader.pages[:max_pages]:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts)
    except Exception:
        return ""


def _read_docx_text(path: Path) -> str:
    """Extract paragraph and table text from a .docx file."""
    try:
        from docx import Document
        doc = Document(str(path))
        chunks = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        chunks.append(cell.text)
        return "\n".join(chunks)
    except Exception:
        return ""


@functools.lru_cache(maxsize=1024)
def _read_pdf_text_cached(path_str: str, mtime: float, size: int) -> str:
    # mtime + size in the cache key invalidate the cache when the file changes
    return _read_pdf_text(Path(path_str))


@functools.lru_cache(maxsize=1024)
def _read_docx_text_cached(path_str: str, mtime: float, size: int) -> str:
    return _read_docx_text(Path(path_str))


def _aggressive_strip(s: str) -> str:
    return re.sub(r"[\s\x00-\x1f]+", "", s)


def _drop_one_variants(token: str):
    """Yield variants of `token` with one internal character removed.

    Used as a fuzzy fallback for PDFs whose font lacks unicode mappings for
    certain glyphs (commonly 'i') — those characters are silently dropped by
    text extractors, so 'education' becomes 'educaton' in the extracted text.
    """
    if len(token) < 6:
        return
    for i in range(1, len(token) - 1):
        yield token[:i] + token[i + 1:]


def cv_signals(text: str) -> tuple[list[str], list[str]]:
    """Return (strong_hits, weak_hits) for CV detection in `text`.

    Exact match first; fuzzy fallback runs only if the exact match wasn't
    decisive (< 1 strong and < 2 weak hits). Fuzzy matches are tagged with
    a trailing '(~)' so diagnostic output makes the distinction visible.
    """
    if not text:
        return [], []
    text = _maybe_despace(text)
    normalized = _normalize(text)

    strong = [kw for kw in _CV_STRONG if kw in normalized]
    weak = [kw for kw in _CV_WEAK if kw in normalized]
    if strong or len(weak) >= 2:
        return strong, weak

    compact = _aggressive_strip(normalized)
    if len(compact) < 50:
        return strong, weak

    for kw in _CV_STRONG:
        if kw in strong:
            continue
        for variant in _drop_one_variants(_aggressive_strip(kw)):
            if variant in compact:
                strong.append(kw + " (~)")
                break
    for kw in _CV_WEAK:
        if kw in weak:
            continue
        for variant in _drop_one_variants(_aggressive_strip(kw)):
            if variant in compact:
                weak.append(kw + " (~)")
                break
    return strong, weak


# ---------------------------------------------------------------------------
# Filename-pattern rules
#
# All match against the normalized filename (see _normalize). The CamelCase
# CV rule is the exception — it matches against the original name so it can
# rely on case transitions ('AcmeCV_', 'JohnCv_2025.pdf' etc.).
# ---------------------------------------------------------------------------

_LB = r"(?<![a-z0-9])"   # left boundary: no preceding letter or digit
_RB = r"(?![a-z0-9])"    # right boundary: no following letter or digit

_CAMELCASE_CV = re.compile(r"(?<=[a-z])(CV|Cv)(?=[_\-\s.]|$)")

NAME_RULES: list[tuple[re.Pattern, str]] = [
    (re.compile(_LB + r"(cv|resume|ozgecmis)" + _RB), "cv"),
    (re.compile(_LB + r"(fatura|invoice|receipt|fis|makbuz|bill)" + _RB),
     "invoices"),
    (re.compile(
        r"(screenshot|screen[\s_-]?shot|ekran[\s_-]?goruntusu|screencap)"
    ), "screenshots"),
    (re.compile(r"^(setup|installer|install)[\s_\-.]"), "installers"),
    (re.compile(
        r"(bordro|gehaltsabrechnung|lohnabrechnung|turnusabrechnung|"
        r"payslip|payroll)"
    ), "payroll"),
    (re.compile(
        r"(dekont|kontoauszug|"
        r"payment[\s_-]+confirmation|"
        r"transaction[\s_-]+details|"
        r"account[\s_-]+statement|"
        r"risk[\s_-]?merkez|"
        r"varlik[\s_-]?degisim)"
    ), "bank"),
    (re.compile(
        r"(steuerberater|mandanteninformation|finanzamt|steuererklarung|"
        r"vergi[\s_-]+(?:levha|beyanname|beyani|dairesi))"
    ), "tax"),
    (re.compile(
        r"(vodafone|turkcell|turk[\s_-]?telekom|numara[\s_-]?tasima)"
    ), "telecom"),
    # Insurance must come before visa: 'ipv antrag' belongs here.
    (re.compile(
        r"(sigorta|versicherung|sfr[\s_-]+ausland|ipv[\s_-]?antrag)"
    ), "insurance"),
    (re.compile(
        r"(sozlesme|mietvertrag|vertragsbestatigung|zusatzvereinbarung|"
        r"kundigung|vollmacht|ibraname|mitgliedschaft|ek[\s_-]+protokol)"
    ), "contracts"),
    (re.compile(
        r"(mietspiegel|yapi[\s_-]?raporu|zemin[\s_-]?yapi|"
        r"nebenkostenabrechnung|emlak|"
        r"(?<![a-z0-9])tapu(?![a-z0-9]))"
    ), "housing"),
    (re.compile(r"(bilet|ticket|boarding[\s_-]?pass)"), "tickets"),
    (re.compile(
        r"(visum|einladungsschreiben|antragszusammenfassung)"
    ), "visa"),
    (re.compile(
        r"(dijital[\s_-]?kimlik|nvi[\s_-]|emniyet[\s_-]|ehliyet|pasaport|"
        r"mezun[\s_-]?belgesi|oturum[\s_-]?uzat|sicil[\s_-]?kayd|"
        r"e[\s_-]?devlet)"
    ), "official"),
    (re.compile(r"(protokoll[\s_-]+theorie|ergebnisprotokoll)"), "exams"),
    (re.compile(
        r"(kilavuz|user[\s_-]?guide|handbuch|"
        r"(?<![a-z0-9])manual(?![a-z0-9]))"
    ), "manuals"),
    (re.compile(
        r"(?<![a-z0-9])(?:iade|refund)(?![a-z0-9])|return[\s_-]+label"
    ), "returns"),
    (re.compile(
        r"(eventlog|errorlog|crashlog|debuglog|"
        r"event[\s_-]+log|error[\s_-]+log|crash[\s_-]+log)"
    ), "logs"),
    (re.compile(
        r"(pdf-?expose-|piaggio|\blimousine\b|"
        r"(?:^|[\s_\-.])(?:bmw|mercedes|audi|vw|volkswagen|renault)[\s_\-])"
    ), "vehicles"),
]


# Extension-to-category mappings come from resources/data/extensions.json.
EXT_RULES = _load_extensions()


# Names skipped during folder scans — sourced from resources/data/skip_names.json.
SKIP_NAMES = _load_skip_names() | {UNDO_LOG_NAME}


def classify_detailed(path: Path, inspect_content: bool = True
                      ) -> tuple[str, str]:
    """Return (category_key, reason_text) for a single file path.

    The reason is a short human-readable note about why the file landed in
    that category — used by the UI when diagnostic mode is on.
    """
    name = path.name

    m = _CAMELCASE_CV.search(name)
    if m:
        return "cv", i18n.t("reason_camelcase_cv", m=m.group())

    name_norm = _normalize(name)
    for pattern, key in NAME_RULES:
        m = pattern.search(name_norm)
        if m:
            return key, i18n.t("reason_name", m=m.group())

    ext = path.suffix.lower()

    if inspect_content and ext in (".pdf", ".docx"):
        try:
            stat = path.stat()
            if ext == ".pdf":
                text = _read_pdf_text_cached(
                    str(path), stat.st_mtime, stat.st_size)
                src = "PDF"
            else:
                text = _read_docx_text_cached(
                    str(path), stat.st_mtime, stat.st_size)
                src = "DOCX"
        except OSError:
            text, src = "", ext.upper().lstrip(".")

        strong, weak = cv_signals(text)
        if strong:
            return "cv", i18n.t(
                "reason_content_strong", src=src, kws=", ".join(strong[:3]))
        if len(weak) >= 2:
            return "cv", i18n.t(
                "reason_content_weak", src=src, n=len(weak),
                kws=", ".join(weak[:4]))

        # Not a CV — produce an explanatory note then fall through to ext.
        if not text:
            note = i18n.t("reason_content_no_text", src=src)
        elif weak:
            note = i18n.t(
                "reason_content_not_cv_weak", src=src, n=len(weak),
                kws=", ".join(weak[:3]))
        else:
            note = i18n.t("reason_content_not_cv", src=src, n=len(text))

        for key, exts in EXT_RULES.items():
            if ext in exts:
                return key, i18n.t(
                    "reason_ext_with_content", ext=ext, note=note)
        return "other", i18n.t("reason_no_match_with_note", note=note)

    for key, exts in EXT_RULES.items():
        if ext in exts:
            return key, i18n.t("reason_ext", ext=ext)
    return "other", i18n.t("reason_ext_no_match", ext=ext)


def classify(path: Path, inspect_content: bool = True) -> str:
    """Convenience wrapper that returns only the category key."""
    return classify_detailed(path, inspect_content)[0]


def plan_moves(root: Path, progress_cb=None, with_reason: bool = False,
               destination: Path | None = None) -> list:
    """Walk `root` (top level only) and produce a move plan.

    Returns a list of tuples. Each tuple is (src, dst, category_key) or
    (src, dst, category_key, reason) when with_reason is True.

    If `destination` is given, every file targets `destination/<Category>/...`
    instead of `root/<Category>/...`. Useful for a centralized library
    workflow where many source folders feed a single organized vault.
    """
    skip_dirs = all_known_category_names()
    target_base = destination if destination else root

    entries = []
    for entry in root.iterdir():
        if entry.is_dir():
            continue
        name = entry.name
        if name in SKIP_NAMES or name.startswith("."):
            continue
        # Defensive: in case a category folder ever shows up as a "file".
        if name in skip_dirs:
            continue
        entries.append(entry)

    moves = []
    total = len(entries)
    for i, entry in enumerate(entries, 1):
        if progress_cb:
            progress_cb(i, total, entry.name)
        key, reason = classify_detailed(entry)
        dst = target_base / category_display(key) / entry.name
        if with_reason:
            moves.append((entry, dst, key, reason))
        else:
            moves.append((entry, dst, key))
    return moves


def resolve_conflict(dst: Path) -> Path:
    """If `dst` exists, return a sibling path with a (1), (2), ... suffix."""
    if not dst.exists():
        return dst
    stem, suffix, parent = dst.stem, dst.suffix, dst.parent
    i = 1
    while True:
        candidate = parent / f"{stem} ({i}){suffix}"
        if not candidate.exists():
            return candidate
        i += 1


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

class SettingsDialog(tk.Toplevel):
    """Modal preferences dialog: language, updates, destination, scheduler."""

    def __init__(self, master: tk.Misc, settings: dict,
                 on_language_change, on_save) -> None:
        super().__init__(master)
        self.transient(master)
        self.resizable(False, False)
        self.title(i18n.t("settings_title"))
        self._on_language_change = on_language_change
        self._on_save = on_save
        self._current_lang = i18n.lang

        frame = ttk.Frame(self, padding=16)
        frame.pack(fill="both", expand=True)
        row = 0

        # --- Language -------------------------------------------------------
        ttk.Label(frame, text=i18n.t("settings_language")).grid(
            row=row, column=0, sticky="w", padx=(0, 10), pady=4)
        self._lang_var = tk.StringVar(value=LANGUAGES[self._current_lang])
        ttk.Combobox(
            frame, state="readonly",
            values=list(LANGUAGES.values()),
            textvariable=self._lang_var,
            width=22,
        ).grid(row=row, column=1, columnspan=2, sticky="w", pady=4)
        row += 1

        # --- Update check ---------------------------------------------------
        self._check_updates_var = tk.BooleanVar(
            value=settings.get("check_updates_on_startup", True))
        ttk.Checkbutton(
            frame, text=i18n.t("settings_check_updates"),
            variable=self._check_updates_var,
        ).grid(row=row, column=0, columnspan=3, sticky="w", pady=(10, 4))
        row += 1

        ttk.Separator(frame, orient="horizontal").grid(
            row=row, column=0, columnspan=3, sticky="ew", pady=(12, 8))
        row += 1

        # --- Destination folder --------------------------------------------
        ttk.Label(frame, text=i18n.t("settings_destination")).grid(
            row=row, column=0, columnspan=3, sticky="w")
        row += 1
        self._dest_var = tk.StringVar(
            value=settings.get("destination_folder", ""))
        ttk.Entry(frame, textvariable=self._dest_var, width=42).grid(
            row=row, column=0, columnspan=2, sticky="we", pady=4)
        ttk.Button(
            frame, text=i18n.t("settings_destination_browse"),
            command=self._pick_destination,
        ).grid(row=row, column=2, sticky="w", padx=(6, 0))
        row += 1
        ttk.Button(
            frame, text=i18n.t("settings_destination_clear"),
            command=lambda: self._dest_var.set(""),
        ).grid(row=row, column=0, sticky="w", pady=(0, 4))
        row += 1
        ttk.Label(frame, text=i18n.t("settings_destination_hint"),
                  foreground="#666", wraplength=420, justify="left").grid(
            row=row, column=0, columnspan=3, sticky="w", pady=(2, 8))
        row += 1

        ttk.Separator(frame, orient="horizontal").grid(
            row=row, column=0, columnspan=3, sticky="ew", pady=(4, 8))
        row += 1

        # --- Background / auto-organize ------------------------------------
        ttk.Label(frame, text=i18n.t("settings_auto_section"),
                  font=("Segoe UI", 10, "bold")).grid(
            row=row, column=0, columnspan=3, sticky="w")
        row += 1
        self._auto_var = tk.BooleanVar(
            value=settings.get("auto_organize", False))
        ttk.Checkbutton(
            frame, text=i18n.t("settings_auto_enable"),
            variable=self._auto_var,
        ).grid(row=row, column=0, columnspan=3, sticky="w", pady=(2, 4))
        row += 1

        ttk.Label(frame, text=i18n.t("settings_auto_folder")).grid(
            row=row, column=0, sticky="w", padx=(0, 8))
        self._auto_folder_var = tk.StringVar(
            value=settings.get("auto_organize_folder", ""))
        ttk.Entry(
            frame, textvariable=self._auto_folder_var, width=32,
        ).grid(row=row, column=1, sticky="we", pady=2)
        ttk.Button(
            frame, text=i18n.t("settings_destination_browse"),
            command=self._pick_auto_folder,
        ).grid(row=row, column=2, sticky="w", padx=(6, 0))
        row += 1

        ttk.Label(frame, text=i18n.t("settings_auto_interval")).grid(
            row=row, column=0, sticky="w", padx=(0, 8))
        self._auto_interval_var = tk.IntVar(
            value=int(settings.get("auto_organize_interval_minutes", 30)))
        ttk.Spinbox(
            frame, from_=MIN_AUTO_INTERVAL_MIN, to=1440,
            textvariable=self._auto_interval_var, width=8,
        ).grid(row=row, column=1, sticky="w", pady=2)
        row += 1

        self._start_minimized_var = tk.BooleanVar(
            value=settings.get("start_in_tray", False))
        ttk.Checkbutton(
            frame, text=i18n.t("settings_auto_start_minimized"),
            variable=self._start_minimized_var,
        ).grid(row=row, column=0, columnspan=3, sticky="w", pady=(4, 4))
        row += 1

        ttk.Label(frame, text=i18n.t("settings_auto_hint"),
                  foreground="#666", wraplength=420, justify="left").grid(
            row=row, column=0, columnspan=3, sticky="w", pady=(2, 8))
        row += 1

        ttk.Label(frame, text=i18n.t("settings_restart_note"),
                  foreground="#666").grid(
            row=row, column=0, columnspan=3, sticky="w", pady=(8, 0))
        row += 1

        btn_row = ttk.Frame(frame)
        btn_row.grid(row=row, column=0, columnspan=3, sticky="e",
                     pady=(16, 0))
        ttk.Button(btn_row, text=i18n.t("settings_cancel"),
                   command=self.destroy).pack(side="right", padx=(8, 0))
        ttk.Button(btn_row, text=i18n.t("settings_save"),
                   command=self._save).pack(side="right")

        frame.columnconfigure(1, weight=1)

        self.grab_set()
        self.wait_visibility()
        self.focus_set()

    def _pick_destination(self) -> None:
        folder = filedialog.askdirectory()
        if folder:
            self._dest_var.set(folder)

    def _pick_auto_folder(self) -> None:
        folder = filedialog.askdirectory()
        if folder:
            self._auto_folder_var.set(folder)

    def _save(self) -> None:
        chosen_label = self._lang_var.get()
        code = next(
            (c for c, label in LANGUAGES.items() if label == chosen_label),
            self._current_lang,
        )
        try:
            interval = max(MIN_AUTO_INTERVAL_MIN,
                           int(self._auto_interval_var.get()))
        except (TypeError, ValueError):
            interval = 30
        self._on_save({
            "language": code,
            "check_updates_on_startup": self._check_updates_var.get(),
            "destination_folder": self._dest_var.get().strip().strip('"'),
            "auto_organize": self._auto_var.get(),
            "auto_organize_folder":
                self._auto_folder_var.get().strip().strip('"'),
            "auto_organize_interval_minutes": interval,
            "start_in_tray": self._start_minimized_var.get(),
        })
        if code != self._current_lang:
            self._on_language_change(code)
        self.destroy()


class OrganizerApp:
    """The top-level Tk window and event handlers for the organizer."""

    def __init__(self, root: tk.Tk, settings: dict) -> None:
        self.root = root
        self.settings = settings
        i18n.set_language(settings.get("language", DEFAULT_LANG))

        self.folder_var = tk.StringVar()
        self.verbose_var = tk.BooleanVar(value=False)
        self.status_var = tk.StringVar()
        self.msg_queue: "queue.Queue[tuple[str, object]]" = queue.Queue()

        # Reactive widget text — updated when language changes.
        self._labels: dict[str, tk.StringVar] = {}

        # Preview plan kept on the instance so right-click can modify it and
        # later organize uses the (possibly modified) plan.
        self._current_plan: list[dict] | None = None

        # Background / tray state. Only activated when auto-organize is on.
        self._tray = None
        self._tray_thread: threading.Thread | None = None
        self._scheduler_stop = threading.Event()
        self._scheduler_pause = threading.Event()
        self._scheduler_thread: threading.Thread | None = None
        self._minimize_to_tray = False  # set when auto-organize is on

        self._build_menu()
        self._build_ui()
        self._apply_language()
        self._bind_shortcuts()
        self._wire_drag_drop()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close_window)

        self.root.after(100, self._drain_queue)

        # Auto-check for updates on launch (background thread, opt-out via
        # Settings). Skipped in development since updating only makes sense
        # for the bundled exe.
        if (self.settings.get("check_updates_on_startup", True)
                and is_running_frozen()):
            self.root.after(800, lambda: self._spawn_update_check(silent=True))

        # Start background mode if configured.
        if self.settings.get("auto_organize", False):
            self._start_background_services()
            if self.settings.get("start_in_tray", False):
                self.root.after(300, self._hide_to_tray)

    # ------- UI construction ------------------------------------------------

    def _label_var(self, key: str) -> tk.StringVar:
        if key not in self._labels:
            self._labels[key] = tk.StringVar(value=i18n.t(key))
        return self._labels[key]

    def _build_menu(self) -> None:
        bar = tk.Menu(self.root)
        self.root.config(menu=bar)

        self._menu_file = tk.Menu(bar, tearoff=0)
        self._menu_settings = tk.Menu(bar, tearoff=0)
        self._menu_help = tk.Menu(bar, tearoff=0)

        bar.add_cascade(menu=self._menu_file, label=i18n.t("menu_file"))
        bar.add_cascade(menu=self._menu_settings,
                        label=i18n.t("menu_settings"))
        bar.add_cascade(menu=self._menu_help, label=i18n.t("menu_help"))

        self._menu_file.add_command(
            label=i18n.t("menu_exit"), command=self.root.destroy)
        self._menu_settings.add_command(
            label=i18n.t("menu_preferences"), command=self._open_settings)
        self._menu_help.add_command(
            label=i18n.t("menu_check_updates"),
            command=self._check_updates_manual)
        self._menu_help.add_separator()
        self._menu_help.add_command(
            label=i18n.t("menu_about"), command=self._show_about)

        self._menubar = bar

    def _build_ui(self) -> None:
        self.root.title(i18n.t("app_title"))
        self.root.geometry("860x620")
        self.root.minsize(700, 500)

        pad = {"padx": 10, "pady": 6}

        # Update banner — only visible when a new release is detected.
        self._update_banner = tk.Frame(
            self.root, bg="#fff8c5", bd=1, relief="solid",
        )
        self._update_banner_label = tk.Label(
            self._update_banner, text="", bg="#fff8c5", anchor="w",
        )
        self._update_banner_label.pack(
            side="left", padx=10, pady=6, fill="x", expand=True)
        self._update_install_btn = tk.Button(
            self._update_banner, text=i18n.t("update_install_btn"),
            command=self._begin_update_install,
        )
        self._update_install_btn.pack(side="left", padx=(0, 6), pady=4)
        self._update_dismiss_btn = tk.Button(
            self._update_banner, text=i18n.t("update_dismiss_btn"),
            command=self._dismiss_update,
        )
        self._update_dismiss_btn.pack(side="left", padx=(0, 8), pady=4)
        # Hidden initially.
        self._pending_update: dict | None = None

        self._top_frame = ttk.Frame(self.root)
        self._top_frame.pack(fill="x", **pad)
        top = self._top_frame
        ttk.Label(top, textvariable=self._label_var("folder_label")).pack(
            side="left")
        self._folder_entry = ttk.Entry(top, textvariable=self.folder_var)
        self._folder_entry.pack(
            side="left", fill="x", expand=True, padx=(6, 6))
        self.browse_btn = ttk.Button(
            top, textvariable=self._label_var("browse_btn"),
            command=self._browse)
        self.browse_btn.pack(side="left")
        self._open_explorer_btn = ttk.Button(
            top, textvariable=self._label_var("open_in_explorer_btn"),
            command=self._open_in_explorer)
        self._open_explorer_btn.pack(side="left", padx=(6, 0))

        # Recent folders row — populated from settings.
        recent_row = ttk.Frame(self.root)
        recent_row.pack(fill="x", padx=10, pady=(0, 4))
        ttk.Label(recent_row,
                  textvariable=self._label_var("recent_folders_label")).pack(
            side="left")
        self.recent_var = tk.StringVar()
        self.recent_combo = ttk.Combobox(
            recent_row, state="readonly", textvariable=self.recent_var,
            values=self.settings.get("recent_folders", []),
        )
        self.recent_combo.pack(side="left", fill="x", expand=True, padx=(6, 0))
        self.recent_combo.bind("<<ComboboxSelected>>", self._on_recent_picked)

        actions = ttk.Frame(self.root)
        actions.pack(fill="x", **pad)
        self.preview_btn = ttk.Button(
            actions, textvariable=self._label_var("preview_btn"),
            command=self._preview)
        self.preview_btn.pack(side="left", padx=(0, 6))
        self.organize_btn = ttk.Button(
            actions, textvariable=self._label_var("organize_btn"),
            command=self._organize)
        self.organize_btn.pack(side="left", padx=(0, 6))
        self.undo_btn = ttk.Button(
            actions, textvariable=self._label_var("undo_btn"),
            command=self._undo)
        self.undo_btn.pack(side="left", padx=(0, 6))
        ttk.Checkbutton(
            actions, textvariable=self._label_var("verbose_check"),
            variable=self.verbose_var).pack(side="left", padx=(12, 0))
        ttk.Button(
            actions, textvariable=self._label_var("clear_log_btn"),
            command=self._clear_log).pack(side="right")

        self.progress = ttk.Progressbar(self.root, mode="determinate")
        self.progress.pack(fill="x", **pad)

        self.status_var.set(i18n.t("status_ready"))
        ttk.Label(self.root, textvariable=self.status_var, anchor="w").pack(
            fill="x", padx=10)

        log_frame = ttk.Frame(self.root)
        log_frame.pack(fill="both", expand=True, **pad)
        self.log = tk.Text(log_frame, wrap="none", font=("Consolas", 9))
        yscroll = ttk.Scrollbar(log_frame, orient="vertical",
                                command=self.log.yview)
        xscroll = ttk.Scrollbar(log_frame, orient="horizontal",
                                command=self.log.xview)
        self.log.configure(yscrollcommand=yscroll.set,
                           xscrollcommand=xscroll.set)
        self.log.grid(row=0, column=0, sticky="nsew")
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll.grid(row=1, column=0, sticky="ew")
        log_frame.rowconfigure(0, weight=1)
        log_frame.columnconfigure(0, weight=1)
        # Right-click on preview lines to reclassify a file.
        self.log.bind("<Button-3>", self._on_preview_right_click)

        footer = ttk.Frame(self.root)
        footer.pack(fill="x", padx=10, pady=(0, 8))
        self.bmc_link = tk.Label(
            footer, text="", fg="#1f6feb", cursor="hand2",
            font=("Segoe UI", 9, "underline"),
        )
        self.bmc_link.pack(side="right")
        self.bmc_link.bind("<Button-1>", lambda _e: webbrowser.open(BMC_URL))

    def _apply_language(self) -> None:
        """Refresh every translatable label after a language change."""
        self.root.title(i18n.t("app_title"))
        for key, var in self._labels.items():
            var.set(i18n.t(key))
        self.status_var.set(i18n.t("status_ready"))

        # Menu cascade labels do not bind to StringVars cleanly across
        # platforms — easier to rebuild the whole menu from scratch.
        self._build_menu()

        # Footer link uses the localized phrasing + URL.
        self.bmc_link.config(text=f"{i18n.t('bmc_label')}  -  {BMC_URL}")

    # ------- Menu actions ---------------------------------------------------

    def _open_settings(self) -> None:
        SettingsDialog(
            self.root,
            settings=self.settings,
            on_language_change=self._change_language,
            on_save=self._save_settings_dict,
        )

    def _save_settings_dict(self, updates: dict) -> None:
        was_auto = self.settings.get("auto_organize", False)
        self.settings.update(updates)
        save_settings(self.settings)
        now_auto = self.settings.get("auto_organize", False)
        if now_auto and not was_auto:
            self._start_background_services()
        elif was_auto and not now_auto:
            self._stop_background_services()

    # ------- Stats dialog ---------------------------------------------------

    def _show_stats_dialog(self, stats: dict) -> None:
        size_str = self._human_bytes(stats.get("bytes", 0))
        body = i18n.t(
            "stats_body",
            total=stats["total"], cats=stats["cats"],
            size=size_str,
            elapsed=f"{stats['elapsed']:.1f}",
            errors=stats["errors"],
        )

        dlg = tk.Toplevel(self.root)
        dlg.title(i18n.t("stats_title"))
        dlg.transient(self.root)
        dlg.resizable(False, False)

        frame = ttk.Frame(dlg, padding=18)
        frame.pack()
        ttk.Label(frame, text=i18n.t("stats_title"),
                  font=("Segoe UI", 11, "bold")).pack(anchor="w")
        ttk.Label(frame, text=body, justify="left").pack(
            anchor="w", pady=(8, 12))

        btn_row = ttk.Frame(frame)
        btn_row.pack(anchor="e")
        folder = stats.get("folder") or ""
        if folder:
            ttk.Button(
                btn_row, text=i18n.t("stats_open_folder"),
                command=lambda: (os.startfile(folder), dlg.destroy()),
            ).pack(side="right", padx=(8, 0))
        ttk.Button(btn_row, text=i18n.t("stats_close"),
                   command=dlg.destroy).pack(side="right")

        dlg.grab_set()
        dlg.focus_set()

    @staticmethod
    def _human_bytes(n: int) -> str:
        size = float(n)
        for unit in ("B", "KB", "MB", "GB", "TB"):
            if size < 1024:
                return f"{size:.1f} {unit}" if unit != "B" else f"{int(size)} B"
            size /= 1024
        return f"{size:.1f} PB"

    # ------- Background services (tray + scheduler) ------------------------

    def _start_background_services(self) -> None:
        self._minimize_to_tray = True
        self._scheduler_stop.clear()
        self._scheduler_pause.clear()
        if not self._scheduler_thread or not self._scheduler_thread.is_alive():
            self._scheduler_thread = threading.Thread(
                target=self._scheduler_loop, daemon=True)
            self._scheduler_thread.start()
        self._ensure_tray_running()

    def _stop_background_services(self) -> None:
        self._minimize_to_tray = False
        self._scheduler_stop.set()
        self._stop_tray()

    def _scheduler_loop(self) -> None:
        """Periodically run organize on the watched folder."""
        # Run once shortly after activation so the user sees activity.
        first_delay = 5
        for _ in range(first_delay):
            if self._scheduler_stop.is_set():
                return
            time.sleep(1)

        while not self._scheduler_stop.is_set():
            interval_min = max(
                MIN_AUTO_INTERVAL_MIN,
                int(self.settings.get("auto_organize_interval_minutes", 30)),
            )
            if not self._scheduler_pause.is_set():
                self._run_scheduled_pass()
            # Sleep in 1s slices so we react to stop/pause quickly.
            for _ in range(interval_min * 60):
                if self._scheduler_stop.is_set():
                    return
                time.sleep(1)

    def _run_scheduled_pass(self) -> None:
        folder = self.settings.get("auto_organize_folder", "").strip()
        if not folder:
            self.msg_queue.put(("log", i18n.t("auto_no_folder_set")))
            return
        root = Path(folder)
        if not root.is_dir():
            return
        self._update_tray_status("running")
        try:
            self._organize_worker(root, silent=True)
            # Notify count via msg_queue → drained on main thread.
            self.msg_queue.put(("tray_notify_done", root))
        finally:
            self._update_tray_status("idle")

    # ----- Tray icon plumbing ------------------------------------------------

    def _ensure_tray_running(self) -> None:
        if self._tray is not None:
            return
        try:
            import pystray
            from PIL import Image, ImageDraw
        except Exception:
            return  # tray libs missing — silently skip

        icon_image = self._make_tray_icon(Image, ImageDraw)
        menu = pystray.Menu(
            pystray.MenuItem(
                i18n.t("tray_show"), self._tray_show, default=True),
            pystray.MenuItem(
                i18n.t("tray_run_now"), self._tray_run_now),
            pystray.MenuItem(
                lambda _: (i18n.t("tray_resume") if
                           self._scheduler_pause.is_set()
                           else i18n.t("tray_pause")),
                self._tray_toggle_pause),
            pystray.Menu.SEPARATOR,
            pystray.MenuItem(i18n.t("tray_quit"), self._tray_quit),
        )
        self._tray = pystray.Icon(
            "file-organizer", icon_image, APP_NAME, menu)
        self._tray_thread = threading.Thread(
            target=self._tray.run, daemon=True)
        self._tray_thread.start()

    @staticmethod
    def _make_tray_icon(image_mod, draw_mod):
        """Render a simple 64x64 'FO' icon for the tray."""
        img = image_mod.new("RGB", (64, 64), color=(31, 111, 235))
        draw = draw_mod.Draw(img)
        draw.rectangle((8, 16, 56, 48), fill=(255, 255, 255))
        draw.rectangle((8, 16, 56, 22), fill=(31, 111, 235))
        return img

    def _update_tray_status(self, state: str) -> None:
        if not self._tray:
            return
        key = {
            "running": "tray_status_running",
            "paused": "tray_status_paused",
        }.get(state, "tray_status_idle")
        try:
            self._tray.title = i18n.t(key)
        except Exception:
            pass

    def _stop_tray(self) -> None:
        if self._tray:
            try:
                self._tray.stop()
            except Exception:
                pass
            self._tray = None

    def _hide_to_tray(self) -> None:
        self._ensure_tray_running()
        try:
            self.root.withdraw()
        except tk.TclError:
            pass

    def _tray_show(self, _icon=None, _item=None) -> None:
        self.root.after(0, self._show_window_from_tray)

    def _show_window_from_tray(self) -> None:
        try:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_force()
        except tk.TclError:
            pass

    def _tray_run_now(self, _icon=None, _item=None) -> None:
        threading.Thread(
            target=self._run_scheduled_pass, daemon=True).start()

    def _tray_toggle_pause(self, _icon=None, _item=None) -> None:
        if self._scheduler_pause.is_set():
            self._scheduler_pause.clear()
            self._update_tray_status("idle")
        else:
            self._scheduler_pause.set()
            self._update_tray_status("paused")

    def _tray_quit(self, _icon=None, _item=None) -> None:
        self._scheduler_stop.set()
        self._stop_tray()
        self.root.after(0, self.root.destroy)

    def _change_language(self, code: str) -> None:
        i18n.set_language(code)
        self.settings["language"] = code
        save_settings(self.settings)
        self._apply_language()
        # If the update banner is showing, refresh its text too.
        if self._pending_update:
            self._show_update_banner(self._pending_update)

    def _show_about(self) -> None:
        messagebox.showinfo(
            i18n.t("about_title"),
            i18n.t("about_body", app=APP_NAME, ver=APP_VERSION),
        )

    # ------- Updates --------------------------------------------------------

    def _check_updates_manual(self) -> None:
        self.status_var.set(i18n.t("update_downloading"))
        self._spawn_update_check(silent=False)

    def _spawn_update_check(self, silent: bool) -> None:
        threading.Thread(
            target=self._update_check_worker, args=(silent,), daemon=True,
        ).start()

    def _update_check_worker(self, silent: bool) -> None:
        info = fetch_latest_release_info()
        if not info:
            if not silent:
                self.msg_queue.put(("update_failed", None))
            return
        if is_newer_version(info["version"], APP_VERSION):
            dismissed = self.settings.get("dismissed_version", "")
            if silent and dismissed and not is_newer_version(
                info["version"], dismissed
            ):
                return
            self.msg_queue.put(("update_found", info))
        elif not silent:
            self.msg_queue.put(("update_none", APP_VERSION))

    def _show_update_banner(self, info: dict) -> None:
        self._pending_update = info
        self._update_banner_label.config(
            text=i18n.t("update_available", version=info["version"]))
        self._update_install_btn.config(text=i18n.t("update_install_btn"))
        self._update_dismiss_btn.config(text=i18n.t("update_dismiss_btn"))
        # Insert banner above everything else.
        self._update_banner.pack(
            side="top", fill="x", padx=10, pady=(6, 0),
            before=self._top_frame,
        )

    def _dismiss_update(self) -> None:
        if self._pending_update:
            self.settings["dismissed_version"] = self._pending_update["version"]
            save_settings(self.settings)
        self._pending_update = None
        self._update_banner.pack_forget()

    def _begin_update_install(self) -> None:
        if not self._pending_update:
            return
        info = self._pending_update
        confirmed = messagebox.askyesno(
            i18n.t("update_confirm_title"),
            i18n.t(
                "update_confirm_body",
                ver=info["version"], cur=APP_VERSION,
                size=_human_size(info["size"]),
            ),
        )
        if not confirmed:
            return
        self._set_busy(True)
        self.status_var.set(i18n.t("update_downloading"))
        self.progress.config(mode="determinate", maximum=100, value=0)
        threading.Thread(
            target=self._update_apply_worker, args=(info["url"],),
            daemon=True,
        ).start()

    def _shutdown_for_update(self) -> None:
        """Tear down the app so the batch can swap the exe.

        Must run on the main thread. We stop the scheduler and tray, destroy
        the Tk root (releasing the GUI handles), and then call os._exit to
        guarantee the process is gone within a tick — this is the file lock
        the swap helper has been waiting on.
        """
        self._scheduler_stop.set()
        try:
            self._stop_tray()
        except Exception:
            pass
        try:
            self.root.destroy()
        except Exception:
            pass
        os._exit(0)

    def _update_apply_worker(self, url: str) -> None:
        def progress(done, total):
            pct = int(done * 100 / total) if total else 0
            self.msg_queue.put(("progress", pct))
        try:
            apply_update(url, on_progress=progress)
        except Exception as e:
            self.msg_queue.put(("update_apply_error", str(e)))
            return
        # Hand control back to the main thread; it owns process shutdown.
        # We MUST NOT call sys.exit/os._exit from a worker thread — sys.exit
        # only stops the current thread, leaving the exe locked and the
        # update batch in an infinite retry loop.
        self.msg_queue.put(("update_apply_done", None))

    # ------- Helpers --------------------------------------------------------

    def _browse(self) -> None:
        folder = filedialog.askdirectory()
        if folder:
            self.folder_var.set(folder)
            self._remember_recent(folder)

    def _open_in_explorer(self) -> None:
        path = self.folder_var.get().strip().strip('"')
        if not path:
            return
        try:
            os.startfile(path)
        except OSError:
            messagebox.showerror(
                i18n.t("app_title"),
                i18n.t("folder_not_found", path=path),
            )

    def _on_recent_picked(self, _event=None) -> None:
        chosen = self.recent_var.get()
        if chosen:
            self.folder_var.set(chosen)

    def _remember_recent(self, folder: str) -> None:
        """Push a folder to the top of the recent list (capped, deduped)."""
        recents = list(self.settings.get("recent_folders", []))
        if folder in recents:
            recents.remove(folder)
        recents.insert(0, folder)
        recents = recents[:RECENT_FOLDERS_LIMIT]
        self.settings["recent_folders"] = recents
        save_settings(self.settings)
        self.recent_combo.config(values=recents)

    def _bind_shortcuts(self) -> None:
        self.root.bind("<Control-p>", lambda _e: self._preview())
        self.root.bind("<Control-P>", lambda _e: self._preview())
        self.root.bind("<Control-o>", lambda _e: self._organize())
        self.root.bind("<Control-O>", lambda _e: self._organize())
        self.root.bind("<Control-z>", lambda _e: self._undo())
        self.root.bind("<Control-Z>", lambda _e: self._undo())
        self.root.bind("<F5>", lambda _e: self._preview())

    def _wire_drag_drop(self) -> None:
        """Accept folder drops on the main window if tkinterdnd2 is loaded."""
        try:
            from tkinterdnd2 import DND_FILES
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind("<<Drop>>", self._on_drop)
        except Exception:
            # Not running under a TkinterDnD root, or module missing — skip.
            pass

    def _on_drop(self, event) -> None:
        raw = event.data or ""
        # Tk wraps paths with spaces in {curly braces}.
        path = raw.strip().strip("{}").split("} {")[0].strip("{}")
        if not path:
            return
        p = Path(path)
        # If a file was dropped, use its parent folder.
        target = p if p.is_dir() else p.parent
        self.folder_var.set(str(target))
        self._remember_recent(str(target))

    def _on_close_window(self) -> None:
        """When auto-organize is on, X button hides to tray instead of quit."""
        if self._minimize_to_tray:
            self._hide_to_tray()
        else:
            self.root.destroy()

    def _log_line(self, msg: str) -> None:
        self.log.insert(END, msg + "\n")
        self.log.see(END)

    def _clear_log(self) -> None:
        self.log.delete("1.0", END)

    def _set_busy(self, busy: bool) -> None:
        state = "disabled" if busy else "normal"
        for btn in (self.preview_btn, self.organize_btn, self.undo_btn):
            btn.config(state=state)

    def _resolve_root(self) -> Path | None:
        raw = self.folder_var.get().strip().strip('"')
        if not raw:
            messagebox.showwarning(
                i18n.t("app_title"), i18n.t("pick_folder_first"))
            return None
        root = Path(raw)
        if not root.is_dir():
            messagebox.showerror(
                i18n.t("app_title"),
                i18n.t("folder_not_found", path=raw))
            return None
        return root

    # ------- Preview --------------------------------------------------------

    def _preview(self) -> None:
        root = self._resolve_root()
        if not root:
            return
        self._remember_recent(str(root))
        self._current_plan = None
        self._clear_log()
        self._log_line(i18n.t("preview_header", path=root))
        self._log_line(i18n.t("preview_scanning"))
        self._set_busy(True)
        verbose = self.verbose_var.get()
        threading.Thread(
            target=self._preview_worker, args=(root, verbose),
            daemon=True,
        ).start()

    def _preview_worker(self, root: Path, verbose: bool) -> None:
        def progress(i, total, name):
            self.msg_queue.put(("progress_max", total))
            self.msg_queue.put(("progress", i))
            self.msg_queue.put((
                "status",
                i18n.t("scanning_progress", i=i, total=total, name=name),
            ))
        try:
            destination = self._destination_path()
            moves = plan_moves(
                root,
                progress_cb=progress,
                with_reason=True,
                destination=destination,
            )
            if not moves:
                self.msg_queue.put(("log", i18n.t("no_files")))
                self.msg_queue.put(("done", i18n.t("empty_done")))
                return

            plan = [
                {"src": src, "dst": dst, "key": key, "reason": reason}
                for src, dst, key, reason in moves
            ]
            self.msg_queue.put(("plan_ready", (plan, root, verbose)))
        except Exception as e:
            self.msg_queue.put(("log", i18n.t("fatal_error", err=e)))
            self.msg_queue.put(("done", i18n.t("error_done")))

    def _destination_path(self) -> Path | None:
        dest = self.settings.get("destination_folder", "").strip()
        if not dest:
            return None
        p = Path(dest)
        if p.is_dir():
            return p
        # Try to create it if the user explicitly configured a missing path.
        try:
            p.mkdir(parents=True, exist_ok=True)
            return p
        except OSError:
            return None

    def _render_plan(self, plan: list[dict], verbose: bool) -> None:
        """Repaint the log area from the current plan."""
        self._clear_log()
        grouped: dict[str, list[dict]] = {}
        for entry in plan:
            label = category_display(entry["key"])
            grouped.setdefault(label, []).append(entry)

        for label in sorted(grouped):
            entries = sorted(grouped[label], key=lambda e: e["src"].name)
            self._log_line(f"[{label}]  ({len(entries)})")
            for entry in entries:
                if verbose:
                    self._log_line(
                        f"  - {entry['src'].name}    [{entry['reason']}]")
                else:
                    self._log_line(f"  - {entry['src'].name}")
            self._log_line("")

        self._log_line(
            i18n.t("preview_summary", total=len(plan), cats=len(grouped)))
        if not verbose:
            self._log_line(i18n.t("verbose_hint"))
        self._log_line(i18n.t("drop_hint"))

    # ------- Organize -------------------------------------------------------

    def _organize(self) -> None:
        root = self._resolve_root()
        if not root:
            return
        if not messagebox.askyesno(
            i18n.t("app_title"),
            i18n.t("confirm_organize", path=root),
        ):
            return
        self._remember_recent(str(root))
        self._set_busy(True)
        threading.Thread(
            target=self._organize_worker, args=(root,), daemon=True,
        ).start()

    def _organize_worker(self, root: Path,
                         silent: bool = False) -> None:
        try:
            # Prefer the previewed plan (so user reclassifications are honoured).
            if self._current_plan:
                moves = [
                    (e["src"], e["dst"], e["key"])
                    for e in self._current_plan
                ]
            else:
                destination = self._destination_path()
                moves = plan_moves(root, destination=destination)

            if not silent:
                self.msg_queue.put((
                    "log", i18n.t("organize_header", path=root)))
            if not moves:
                if not silent:
                    self.msg_queue.put(("log", i18n.t("no_files")))
                    self.msg_queue.put(("done", i18n.t("empty_done")))
                return
            self.msg_queue.put(("progress_max", len(moves)))

            undo_records = []
            errors = 0
            total_bytes = 0
            timestamp = datetime.now().isoformat(timespec="seconds")
            started = time.monotonic()

            per_category: dict[str, int] = {}

            for i, (src, dst, key) in enumerate(moves, 1):
                try:
                    size = src.stat().st_size
                except OSError:
                    size = 0
                try:
                    dst.parent.mkdir(parents=True, exist_ok=True)
                    final = resolve_conflict(dst)
                    shutil.move(str(src), str(final))
                    undo_records.append({
                        "from": str(final), "to": str(src),
                    })
                    total_bytes += size
                    per_category[key] = per_category.get(key, 0) + 1
                    if not silent:
                        self.msg_queue.put((
                            "log",
                            f"[{category_display(key)}] {src.name}",
                        ))
                except Exception as e:
                    errors += 1
                    if not silent:
                        self.msg_queue.put((
                            "log",
                            f"{i18n.t('error_label')}: {src.name} -> {e}",
                        ))
                self.msg_queue.put(("progress", i))

            elapsed = time.monotonic() - started
            self._append_undo(root, timestamp, undo_records)
            self._current_plan = None  # reset after applying

            if not silent:
                self.msg_queue.put((
                    "log",
                    i18n.t("organize_done_log",
                           moved=len(undo_records), errors=errors),
                ))
                self.msg_queue.put((
                    "done",
                    i18n.t("organize_done_status",
                           moved=len(undo_records), errors=errors),
                ))
                self.msg_queue.put(("stats", {
                    "folder": str(root),
                    "total": len(undo_records),
                    "cats": len(per_category),
                    "bytes": total_bytes,
                    "elapsed": elapsed,
                    "errors": errors,
                }))
        except Exception as e:
            if not silent:
                self.msg_queue.put(("log", i18n.t("fatal_error", err=e)))
                self.msg_queue.put(("done", i18n.t("error_done")))

    def _append_undo(self, root: Path, timestamp: str,
                     records: list[dict]) -> None:
        undo_path = root / UNDO_LOG_NAME
        history = []
        if undo_path.exists():
            try:
                history = json.loads(undo_path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                history = []
        history.append({"timestamp": timestamp, "moves": records})
        try:
            undo_path.write_text(
                json.dumps(history, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        except OSError as e:
            self.msg_queue.put((
                "log", i18n.t("warn_undo_write_failed", err=e)))

    # ------- Undo -----------------------------------------------------------

    def _undo(self) -> None:
        root = self._resolve_root()
        if not root:
            return
        undo_path = root / UNDO_LOG_NAME
        if not undo_path.exists():
            messagebox.showinfo(
                i18n.t("app_title"), i18n.t("undo_no_history"))
            return
        try:
            history = json.loads(undo_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError) as e:
            messagebox.showerror(
                i18n.t("app_title"),
                i18n.t("undo_read_error", err=e))
            return
        if not history:
            messagebox.showinfo(
                i18n.t("app_title"), i18n.t("undo_no_history"))
            return

        last = history[-1]
        if not messagebox.askyesno(
            i18n.t("app_title"),
            i18n.t("confirm_undo",
                   n=len(last["moves"]), date=last["timestamp"]),
        ):
            return

        self._set_busy(True)
        threading.Thread(
            target=self._undo_worker, args=(root, history),
            daemon=True,
        ).start()

    def _undo_worker(self, root: Path, history: list) -> None:
        try:
            last = history.pop()
            self.msg_queue.put((
                "log", i18n.t("undo_header", date=last["timestamp"])))
            moves = last["moves"]
            self.msg_queue.put(("progress_max", len(moves)))
            errors = 0
            for i, m in enumerate(reversed(moves), 1):
                src, dst = Path(m["from"]), Path(m["to"])
                try:
                    if src.exists():
                        dst.parent.mkdir(parents=True, exist_ok=True)
                        final = resolve_conflict(dst)
                        shutil.move(str(src), str(final))
                        self.msg_queue.put(("log", f"<- {final.name}"))
                    else:
                        self.msg_queue.put((
                            "log", i18n.t("skipped_missing", name=src.name)))
                except Exception as e:
                    errors += 1
                    self.msg_queue.put((
                        "log",
                        f"{i18n.t('error_label')}: {src} -> {e}",
                    ))
                self.msg_queue.put(("progress", i))

            # Remove now-empty category folders (any language).
            for label in all_known_category_names():
                folder = root / label
                if folder.is_dir():
                    try:
                        if not any(folder.iterdir()):
                            folder.rmdir()
                    except OSError:
                        pass

            self._save_remaining_undo(root, history)
            self.msg_queue.put((
                "log",
                i18n.t("undo_done_log",
                       moved=len(moves) - errors, errors=errors),
            ))
            self.msg_queue.put(("done", i18n.t("undo_done")))
        except Exception as e:
            self.msg_queue.put(("log", i18n.t("fatal_error", err=e)))
            self.msg_queue.put(("done", i18n.t("error_done")))

    @staticmethod
    def _save_remaining_undo(root: Path, history: list) -> None:
        undo_path = root / UNDO_LOG_NAME
        if history:
            try:
                undo_path.write_text(
                    json.dumps(history, indent=2, ensure_ascii=False),
                    encoding="utf-8",
                )
            except OSError:
                pass
        else:
            try:
                undo_path.unlink()
            except OSError:
                pass

    # ------- Message pump ---------------------------------------------------

    def _drain_queue(self) -> None:
        try:
            while True:
                kind, payload = self.msg_queue.get_nowait()
                if kind == "log":
                    self._log_line(payload)
                elif kind == "progress_max":
                    self.progress.config(maximum=payload, value=0)
                elif kind == "progress":
                    self.progress.config(value=payload)
                elif kind == "status":
                    self.status_var.set(payload)
                elif kind == "done":
                    self.status_var.set(payload)
                    self._set_busy(False)
                elif kind == "update_found":
                    self._show_update_banner(payload)
                elif kind == "update_none":
                    messagebox.showinfo(
                        i18n.t("app_title"),
                        i18n.t("update_no_update", version=payload),
                    )
                    self.status_var.set(i18n.t("status_ready"))
                elif kind == "update_failed":
                    messagebox.showwarning(
                        i18n.t("app_title"),
                        i18n.t("update_check_failed"),
                    )
                    self.status_var.set(i18n.t("status_ready"))
                elif kind == "update_apply_error":
                    self._set_busy(False)
                    self.progress.config(value=0)
                    self.status_var.set(i18n.t("status_ready"))
                    messagebox.showerror(
                        i18n.t("app_title"),
                        i18n.t("update_apply_failed", err=payload),
                    )
                elif kind == "update_apply_done":
                    self._shutdown_for_update()
                elif kind == "plan_ready":
                    plan, _root, verbose = payload
                    self._current_plan = plan
                    self._current_verbose = verbose
                    self._render_plan(plan, verbose)
                    self.status_var.set(
                        i18n.t("preview_done", n=len(plan)))
                    self._set_busy(False)
                elif kind == "stats":
                    self._show_stats_dialog(payload)
                elif kind == "tray_notify_done":
                    folder = payload
                    msg = i18n.t("tray_notify_done", n="?", folder=folder)
                    if self._tray:
                        try:
                            self._tray.notify(msg, APP_NAME)
                        except Exception:
                            pass
        except queue.Empty:
            pass
        self.root.after(100, self._drain_queue)

    # ------- Manual reclassification ----------------------------------------

    def _on_preview_right_click(self, event) -> None:
        """Show a category-picker menu when a file line is right-clicked."""
        if not self._current_plan:
            return
        index = self.log.index(f"@{event.x},{event.y}")
        line_text = self.log.get(f"{index} linestart", f"{index} lineend")
        match = re.match(r"^\s*-\s+(.+?)(?:\s{2,}\[.*\])?$", line_text)
        if not match:
            return
        name = match.group(1).strip()
        entry = next(
            (e for e in self._current_plan if e["src"].name == name), None)
        if not entry:
            return

        menu = tk.Menu(self.root, tearoff=0)
        # Sorted display names with key payload so we can update entry["key"].
        active_lang = i18n.lang
        keys_sorted = sorted(
            CATEGORY_NAMES[active_lang].keys(),
            key=lambda k: category_display(k),
        )
        for key in keys_sorted:
            label = category_display(key)
            current_marker = "  ◉" if key == entry["key"] else "   "
            menu.add_command(
                label=f"{current_marker} {label}",
                command=lambda k=key: self._reassign(entry, k),
            )
        menu.tk_popup(event.x_root, event.y_root)

    def _reassign(self, entry: dict, new_key: str) -> None:
        if new_key == entry["key"]:
            return
        entry["key"] = new_key
        entry["reason"] = "user override"
        # Recompute destination path with the new category.
        base = (self._destination_path() or entry["src"].parent)
        entry["dst"] = base / category_display(new_key) / entry["src"].name
        self._render_plan(self._current_plan, self._current_verbose)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    settings = load_settings()
    # Prefer the TkinterDnD root so drag-and-drop works; fall back to plain Tk
    # if the library is unavailable for any reason.
    try:
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except Exception:
        root = tk.Tk()
    try:
        style = ttk.Style()
        if "vista" in style.theme_names():
            style.theme_use("vista")
    except tk.TclError:
        pass
    OrganizerApp(root, settings)
    root.mainloop()


if __name__ == "__main__":
    main()
