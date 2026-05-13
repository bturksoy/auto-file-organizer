"""PDF/DOCX text extraction with caching, plus CV signal detection."""
from __future__ import annotations

import functools
import json
import logging
from pathlib import Path

from app.core.normalize import aggressive_strip, maybe_despace, normalize
from app.core.utils import resources_dir

logging.getLogger("pypdf").setLevel(logging.ERROR)


_CV_KEYWORDS_CACHE: tuple[tuple[str, ...], tuple[str, ...]] | None = None


def _cv_keywords() -> tuple[tuple[str, ...], tuple[str, ...]]:
    """Lazy-load the CV detection word lists once per process."""
    global _CV_KEYWORDS_CACHE
    if _CV_KEYWORDS_CACHE is not None:
        return _CV_KEYWORDS_CACHE
    try:
        raw = json.loads(
            (resources_dir() / "data" / "cv_keywords.json")
            .read_text(encoding="utf-8")
        )
    except (OSError, json.JSONDecodeError):
        _CV_KEYWORDS_CACHE = ((), ())
        return _CV_KEYWORDS_CACHE
    strong = tuple(normalize(k) for k in raw.get("strong", []))
    weak = tuple(normalize(k) for k in raw.get("weak", []))
    _CV_KEYWORDS_CACHE = (strong, weak)
    return _CV_KEYWORDS_CACHE


def read_pdf_text(path: Path, max_pages: int = 4) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path), strict=False)
        if reader.is_encrypted:
            return ""
        chunks = []
        for page in reader.pages[:max_pages]:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(chunks)
    except Exception:
        return ""


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


@functools.lru_cache(maxsize=1024)
def read_pdf_text_cached(path_str: str, mtime: float, size: int) -> str:
    return read_pdf_text(Path(path_str))


@functools.lru_cache(maxsize=1024)
def read_docx_text_cached(path_str: str, mtime: float, size: int) -> str:
    return read_docx_text(Path(path_str))


def _drop_one_variants(token: str):
    if len(token) < 6:
        return
    for i in range(1, len(token) - 1):
        yield token[:i] + token[i + 1:]


def cv_signals(text: str) -> tuple[list[str], list[str]]:
    """Return (strong_hits, weak_hits). Fuzzy fallback for broken glyph PDFs.

    A file is classified as a CV when there is at least one strong hit or
    two weak hits. Fuzzy hits are appended with a trailing "(~)" so the
    diagnostic UI can distinguish them.
    """
    if not text:
        return [], []
    strong_set, weak_set = _cv_keywords()
    text = maybe_despace(text)
    normalized = normalize(text)
    strong = [kw for kw in strong_set if kw in normalized]
    weak = [kw for kw in weak_set if kw in normalized]
    if strong or len(weak) >= 2:
        return strong, weak

    compact = aggressive_strip(normalized)
    if len(compact) < 50:
        return strong, weak

    for kw in strong_set:
        if kw in strong:
            continue
        for variant in _drop_one_variants(aggressive_strip(kw)):
            if variant in compact:
                strong.append(kw + " (~)")
                break
    for kw in weak_set:
        if kw in weak:
            continue
        for variant in _drop_one_variants(aggressive_strip(kw)):
            if variant in compact:
                weak.append(kw + " (~)")
                break
    return strong, weak


def looks_like_cv(text: str) -> bool:
    strong, weak = cv_signals(text)
    return bool(strong) or len(weak) >= 2
